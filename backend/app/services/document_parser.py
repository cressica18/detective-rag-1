"""Document parser — dispatches by file extension."""
import os
import json
from typing import Optional
from app.utils.logging import get_logger

logger = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def guess_doc_type(filename: str) -> str:
    """Guess document type from filename patterns."""
    name = filename.lower()
    if "police_report" in name or "incident_report" in name:
        return "police_report"
    elif "witness_statement" in name or "statement" in name:
        return "witness_statement"
    elif "interview" in name:
        return "interview"
    elif "cctv" in name or "surveillance" in name or "camera" in name:
        return "cctv_transcript"
    elif "forensic" in name or "lab" in name:
        return "forensic_report"
    elif "autopsy" in name or "medical" in name or "coroner" in name:
        return "autopsy_report"
    elif "phone" in name or "call_log" in name or "telecom" in name:
        return "phone_record"
    elif "security_log" in name or "badge" in name or "access_log" in name:
        return "security_log"
    elif "evidence" in name or "inventory" in name or "exhibit" in name:
        return "evidence_inventory"
    elif "store" in name or "receipt" in name or "payment" in name:
        return "evidence_inventory"
    else:
        return "unknown"


def parse_file(filepath: str) -> list[tuple[int, str]]:
    """Parse a file and return list of (page_number, text) tuples."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext == ".docx":
        return _parse_docx(filepath)
    elif ext == ".txt":
        return _parse_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(filepath: str) -> list[tuple[int, str]]:
    try:
        import fitz  # PyMuPDF
        pages = []
        doc = fitz.open(filepath)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                pages.append((page_num + 1, text))
        doc.close()
        if not pages:
            logger.warning(f"PDF {filepath} returned no text — may be scanned/image-based.")
        return pages
    except Exception as e:
        logger.error(f"PDF parse error {filepath}: {e}")
        return [(1, "")]


def _parse_docx(filepath: str) -> list[tuple[int, str]]:
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        full_text = "\n".join(parts)
        return [(1, full_text)]
    except Exception as e:
        logger.error(f"DOCX parse error {filepath}: {e}")
        return [(1, "")]


def _parse_txt(filepath: str) -> list[tuple[int, str]]:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return [(1, text)]
    except Exception as e:
        logger.error(f"TXT parse error {filepath}: {e}")
        return [(1, "")]


def cache_parsed(processed_path: str, case_id: str, doc_id: str, pages: list[tuple[int, str]]):
    """Cache parsed document to JSON."""
    case_dir = os.path.join(processed_path, case_id)
    os.makedirs(case_dir, exist_ok=True)
    cache_file = os.path.join(case_dir, f"{doc_id}.json")
    with open(cache_file, "w", encoding="utf-8") as f:
        json.dump([{"page": p, "text": t} for p, t in pages], f)


def load_cached(processed_path: str, case_id: str, doc_id: str) -> Optional[list[tuple[int, str]]]:
    """Load cached parsed document if it exists."""
    cache_file = os.path.join(processed_path, case_id, f"{doc_id}.json")
    if not os.path.exists(cache_file):
        return None
    with open(cache_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [(item["page"], item["text"]) for item in data]
